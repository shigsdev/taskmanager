"""Import service — OneNote text parsing + Excel goals parsing.

Two import flows:
1. OneNote tasks: user pastes plain text → parser extracts task lines →
   preview → confirmed tasks land in Inbox
2. Excel goals: user uploads .xlsx → parser reads rows into goal
   candidates → preview → confirmed goals created

Both flows log the operation to ImportLog for audit trail.
Duplicate detection prevents re-importing the same items.
"""
from __future__ import annotations

import io
import json
import logging
import os
import re
import uuid
from datetime import date
from typing import Any

from sqlalchemy import select

from models import (
    Goal,
    GoalCategory,
    GoalPriority,
    GoalStatus,
    ImportLog,
    Project,
    ProjectStatus,
    ProjectType,
    Task,
    TaskStatus,
    TaskType,
    Tier,
    db,
)

logger = logging.getLogger(__name__)


# --- OneNote text parsing ----------------------------------------------------


def parse_onenote_text(text: str) -> list[dict[str, Any]]:
    """Parse pasted OneNote text into task candidates.

    OneNote exports tasks as plain text with various bullet styles:
    - Bullet points (-, *, •)
    - Numbered lists (1., 2., etc.)
    - Checkbox items (☐, [], [ ])
    - Plain lines (non-empty, non-header)

    Lines that look like headers (ALL CAPS, very short, or date-only)
    are skipped. Empty lines are skipped.

    Args:
        text: Raw pasted text from OneNote.

    Returns:
        List of candidate dicts with title, type, included fields.
    """
    if not text or not text.strip():
        return []

    candidates = []
    seen_titles: set[str] = set()

    for line in text.splitlines():
        title = _clean_task_line(line)
        if not title:
            continue

        if _is_header_line(title):
            continue

        # Deduplicate within the same paste
        normalized = title.lower().strip()
        if normalized in seen_titles:
            continue
        seen_titles.add(normalized)

        candidates.append({
            "title": title,
            "type": "work",
            "included": True,
        })

    return candidates


def _clean_task_line(line: str) -> str:
    """Strip bullet markers, checkboxes, numbering from a line."""
    line = line.strip()
    if not line:
        return ""

    # Remove checkbox markers: ☐, ☑, ✓, [x], [ ], []
    line = re.sub(r"^[\u2610\u2611\u2713]\s*", "", line)
    line = re.sub(r"^\[[ xX]?\]\s*", "", line)

    # Remove bullet markers: -, *, •, ‣, ▪
    line = re.sub(r"^[-*\u2022\u2023\u25AA]\s+", "", line)

    # Remove numbered list: 1., 2), etc.
    line = re.sub(r"^\d+[.)]\s+", "", line)

    # Remove leading/trailing whitespace and control chars
    line = line.strip()

    # Skip very short lines (likely artifacts)
    if len(line) < 2:
        return ""

    return line


def _is_header_line(text: str) -> bool:
    """Detect header-like lines that aren't real tasks."""
    # All caps lines (any length, up to 40 chars) are likely section headers
    words = text.split()
    if (
        len(words) >= 1
        and len(text) <= 40
        and text == text.upper()
        and not any(c.isdigit() for c in text)
    ):
        return True

    # Date-only lines (e.g., "April 5, 2026", "2026-04-05")
    if re.match(r"^\d{4}-\d{2}-\d{2}$", text):
        return True
    return bool(re.match(
        r"^(?:January|February|March|April|May|June|July|August|"
        r"September|October|November|December)\s+\d{1,2},?\s*\d{4}$",
        text,
        re.IGNORECASE,
    ))


# --- OneNote .docx file parsing ----------------------------------------------


def parse_onenote_docx(file_bytes: bytes) -> list[dict[str, Any]]:
    """Parse a OneNote-exported .docx file into task candidates.

    OneNote can export pages as Word documents via File → Export → Word.
    The exported file contains paragraphs with the same bullet/checkbox
    formatting as the pasted text. We extract each paragraph's text and
    run it through the same cleaning pipeline as parse_onenote_text.

    Args:
        file_bytes: Raw bytes of the .docx file.

    Returns:
        List of candidate dicts with title, type, included fields.
    """
    import docx

    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Cannot read Word document: {e}") from e

    # Extract all paragraph text, one per line
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Also extract text from tables (OneNote sometimes uses tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    if not lines:
        return []

    # Reuse the same text parser
    combined = "\n".join(lines)
    return parse_onenote_text(combined)


# --- Meeting transcript parsing (HyNote / Notion / generic) ------------------

# Markdown / plain-text headers that mark a pre-extracted action-items
# section. HyNote and Notion AI Meeting Notes both emit something like
# this; matching the section verbatim cuts hallucination + tokens vs.
# re-extracting from the whole transcript.
_ACTION_ITEMS_HEADER_RE = re.compile(
    r"^\s{0,3}(?:#{1,6}\s+|\*\*\s*)?(?:action\s*items?|next\s+steps?|"
    r"to[\-\s]?dos?|todos?|follow[\-\s]?ups?)\s*:?\s*(?:\*\*)?\s*$",
    re.IGNORECASE,
)
# Header that opens any *other* section — used to bound the action-items
# block so we don't sweep up the entire rest of the document.
_OTHER_SECTION_HEADER_RE = re.compile(r"^\s{0,3}#{1,6}\s+\S")


_TRANSCRIPT_PROMPT = """\
You are a task extraction assistant. The input is a meeting note or \
transcript (often exported from HyNote, Notion AI Meeting Notes, or a \
similar tool). Your job is to extract concrete, actionable items the \
user committed to or needs to follow up on.

Rules:
- Each item must be a short imperative phrase under 100 characters
  (e.g. "Email Sarah the Q3 plan", "Schedule design review").
- Drop discussion summary, attendee names, meeting metadata, status
  observations, decisions, and anything that is not an action.
- Do not invent items the speaker didn't actually commit to.
- If the transcript explicitly contains a section labeled "Action
  Items" / "Next Steps" / "To-dos" / "Follow-ups" with bulleted
  entries, prefer those entries verbatim (lightly cleaned). Otherwise
  scan the full content.
- Return ONLY a JSON array of objects, no other text.

Each object has:
- title (string, required)
- notes (string or null) — optional one-line context if useful

Example output:
[{"title": "Email Sarah the Q3 plan", "notes": null},
 {"title": "Schedule design review with Alex", "notes": "before EOM"}]

Transcript:
{transcript}
"""


def extract_action_items_section(text: str) -> str | None:
    """Return the body of an explicit 'Action Items' section if present.

    Walks line-by-line for a header matching ``_ACTION_ITEMS_HEADER_RE``,
    then collects subsequent lines until the next markdown header (or
    end of input). Returns the collected body as a single string, or
    ``None`` if no such section was found.

    Trusting the section verbatim is option (a) from the design spec —
    fewer Claude tokens + lower hallucination than re-extracting from
    the whole transcript every time.
    """
    if not text:
        return None
    lines = text.splitlines()
    body: list[str] | None = None
    for raw in lines:
        if body is None:
            if _ACTION_ITEMS_HEADER_RE.match(raw):
                body = []
            continue
        # Inside the section — stop at the next markdown header.
        if _OTHER_SECTION_HEADER_RE.match(raw):
            break
        body.append(raw)
    if body is None:
        return None
    joined = "\n".join(body).strip()
    return joined or None


def parse_transcript_text(text: str) -> list[dict[str, Any]]:
    """Extract action-item task candidates from a meeting transcript.

    Strategy (option (a) — "trust pre-extracted section if present"):
    1. Look for an explicit "Action Items" / "Next Steps" / "To-dos" /
       "Follow-ups" markdown section. If found, send only that section
       to Claude — fewer tokens, less hallucination room.
    2. Otherwise send the whole transcript with the same prompt; Claude
       still scans for action items end-to-end.

    Output shape mirrors ``parse_onenote_text`` so the existing
    ``/api/import/tasks/confirm`` path can ingest these candidates
    unchanged: ``{title, type, included, notes?}``. Default type=work,
    default tier=Inbox (applied at confirm time).

    Raises:
        RuntimeError: If ``ANTHROPIC_API_KEY`` is missing.
    """
    if not text or not text.strip():
        return []

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY not configured")

    section = extract_action_items_section(text)
    payload = section if section else text
    raw = _call_claude_for_transcript(api_key, payload)

    candidates: list[dict[str, Any]] = []
    seen_titles: set[str] = set()
    for item in raw:
        if not isinstance(item, dict):
            continue
        title = item.get("title")
        if not isinstance(title, str):
            continue
        title = title.strip()
        if not title or len(title) < 2:
            continue
        if len(title) > 100:
            title = title[:100]
        normalized = title.lower()
        if normalized in seen_titles:
            continue
        seen_titles.add(normalized)
        notes_raw = item.get("notes")
        notes = notes_raw.strip() if isinstance(notes_raw, str) else ""
        candidates.append({
            "title": title,
            "type": "work",
            "included": True,
            "notes": notes or "",
        })
    return candidates


def _call_claude_for_transcript(api_key: str, transcript: str) -> list[Any]:
    """Make the Claude API call for transcript action-item extraction.

    Separated for testability — tests patch this instead of the HTTP
    layer. Returns the raw list of objects parsed from Claude's reply
    (caller does the cleaning/coercion).
    """
    from egress import EgressError, safe_call_api

    prompt = _TRANSCRIPT_PROMPT.format(transcript=transcript)
    try:
        data = safe_call_api(
            url="https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            json={
                "model": "claude-sonnet-4-6",
                "max_tokens": 2048,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout_sec=60,
            vendor="Claude",
        )
    except EgressError as e:
        raise RuntimeError(str(e)) from e

    content = data.get("content", [{}])[0].get("text", "")
    return _extract_transcript_json_array(content)


def _extract_transcript_json_array(text: str) -> list[Any]:
    """Extract a JSON array of objects from Claude's reply.

    Mirrors ``scan_service._extract_json_object_list`` — direct parse,
    then markdown code fence, then bracket-bound fallback. Returns an
    empty list on any failure rather than raising, so a Claude format
    blip surfaces as "no candidates found" instead of a 500.
    """
    text = (text or "").strip()
    if not text:
        return []

    # Direct parse.
    try:
        parsed = json.loads(text)
        if isinstance(parsed, list):
            return parsed
    except json.JSONDecodeError:
        pass

    # Markdown code fence.
    if "```" in text:
        for part in text.split("```"):
            cleaned = part.strip()
            if cleaned.startswith("json"):
                cleaned = cleaned[4:].strip()
            try:
                parsed = json.loads(cleaned)
                if isinstance(parsed, list):
                    return parsed
            except json.JSONDecodeError:
                continue

    # Bracket-bound fallback.
    start = text.find("[")
    end = text.rfind("]")
    if start != -1 and end != -1 and end > start:
        try:
            parsed = json.loads(text[start : end + 1])
            if isinstance(parsed, list):
                return parsed
        except json.JSONDecodeError:
            pass
    return []


# --- Excel goals parsing ----------------------------------------------------


def parse_excel_goals(file_bytes: bytes) -> list[dict[str, Any]]:
    """Parse an Excel (.xlsx) file into goal candidates.

    Expected columns (case-insensitive, matched by header row):
    - title (required)
    - category: health, personal_growth, relationships, work
    - priority: must, should, could, need_more_info
    - actions: free text
    - target_quarter: e.g. "Q2 2026"
    - status: not_started, in_progress, done, on_hold
    - notes: free text

    The first row is treated as the header. Rows with no title are skipped.

    Args:
        file_bytes: Raw bytes of the .xlsx file.

    Returns:
        List of goal candidate dicts.
    """
    import openpyxl

    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True)
    except Exception as e:
        raise ValueError(f"Cannot read Excel file: {e}") from e

    ws = wb.active
    if ws is None:
        raise ValueError("Excel file has no active worksheet")

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    # Map header names to column indices
    headers = [str(h).strip().lower() if h else "" for h in rows[0]]
    col_map = {name: idx for idx, name in enumerate(headers) if name}

    if "title" not in col_map:
        raise ValueError(
            "Excel file must have a 'title' column in the first row. "
            f"Found columns: {[h for h in headers if h]}"
        )

    candidates = []
    seen_titles: set[str] = set()

    for row in rows[1:]:
        title = _cell_str(row, col_map.get("title"))
        if not title:
            continue

        # Deduplicate within file
        normalized = title.lower()
        if normalized in seen_titles:
            continue
        seen_titles.add(normalized)

        category = _cell_str(row, col_map.get("category")) or "work"
        priority = _cell_str(row, col_map.get("priority")) or "should"

        # Validate enum values, fall back to defaults
        if not _valid_enum(GoalCategory, category):
            category = "work"
        if not _valid_enum(GoalPriority, priority):
            priority = "should"

        status = _cell_str(row, col_map.get("status")) or "not_started"
        if not _valid_enum(GoalStatus, status):
            status = "not_started"

        candidates.append({
            "title": title,
            "category": category,
            "priority": priority,
            "actions": _cell_str(row, col_map.get("actions")) or "",
            "target_quarter": _cell_str(row, col_map.get("target_quarter")) or "",
            "status": status,
            "notes": _cell_str(row, col_map.get("notes")) or "",
            "included": True,
        })

    return candidates


def _cell_str(row: tuple, idx: int | None) -> str:
    """Safely extract a string from a row tuple by index."""
    if idx is None or idx >= len(row) or row[idx] is None:
        return ""
    return str(row[idx]).strip()


def _valid_enum(enum_cls, value: str) -> bool:
    """Check if a string is a valid member of an enum."""
    try:
        enum_cls(value)
        return True
    except ValueError:
        return False


# --- Duplicate detection -----------------------------------------------------


def find_duplicate_tasks(titles: list[str]) -> list[str]:
    """Return titles that already exist as active tasks (case-insensitive)."""
    if not titles:
        return []

    from sqlalchemy import func

    lower_titles = [t.lower() for t in titles]
    existing = db.session.scalars(
        select(Task.title).where(
            Task.status != TaskStatus.DELETED,
            func.lower(Task.title).in_(lower_titles),
        )
    ).all()
    existing_lower = {t.lower() for t in existing}

    return [t for t in titles if t.lower() in existing_lower]


def parse_project_names_text(text: str) -> list[dict[str, Any]]:
    """#80 (2026-04-26): parse one project name per line from raw text.

    Strips bullet/numbered list markers (- *, 1., etc.) so a copy-pasted
    list works without manual cleanup. Skips empty lines and obvious
    headers. Returns the same candidate shape the UI consumes.
    """
    if not text or not text.strip():
        return []

    candidates: list[dict[str, Any]] = []
    seen_lower: set[str] = set()
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        # Strip bullet / numbered prefixes (mirrors parse_onenote_text logic).
        cleaned = _clean_task_line(line)
        if not cleaned:
            continue
        if _is_header_line(cleaned):
            continue
        if len(cleaned) < 2:
            continue
        lo = cleaned.lower()
        if lo in seen_lower:
            continue
        seen_lower.add(lo)
        candidates.append({
            "name": cleaned,
            "type": "work",
            "included": True,
        })
    return candidates


def parse_excel_tasks(file_bytes: bytes) -> list[dict[str, Any]]:
    """#89 (2026-04-26): parse a .xlsx of task rows.

    Expected header row (case-insensitive):
      - title (required)
      - type: work, personal — default work
      - tier: inbox, today, tomorrow, this_week, next_week, backlog,
        freezer — default inbox
      - due_date: ISO YYYY-MM-DD; openpyxl returns datetime which gets
        normalised to YYYY-MM-DD string
      - linked_goal: free-text goal title; matched case-insensitive at
        create time. No exact-id requirement.
      - linked_project: free-text project name; same matching.
      - notes: free text
      - url: free text

    Rows with no title are skipped. Returns the candidate shape the
    review UI consumes (the field mapping mirrors what
    `create_tasks_from_import` already accepts).
    """
    import openpyxl

    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True)
    except Exception as e:
        raise ValueError(f"Cannot read Excel file: {e}") from e

    ws = wb.active
    if ws is None:
        raise ValueError("Excel file has no active worksheet")

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    headers = [str(h).strip().lower() if h else "" for h in rows[0]]
    col_map = {name: idx for idx, name in enumerate(headers) if name}

    candidates: list[dict[str, Any]] = []
    for row in rows[1:]:
        title = _cell_str(row, col_map.get("title"))
        if not title:
            continue
        type_str = _cell_str(row, col_map.get("type")).lower()
        if not _valid_enum(TaskType, type_str):
            type_str = "work"
        tier_str = _cell_str(row, col_map.get("tier")).lower()
        if not _valid_enum(Tier, tier_str):
            tier_str = "inbox"
        # due_date may come back from openpyxl as a date / datetime —
        # normalise to YYYY-MM-DD string. Plain strings pass through.
        due_idx = col_map.get("due_date")
        due_raw = row[due_idx] if (due_idx is not None and due_idx < len(row)) else None
        if hasattr(due_raw, "strftime"):
            due_str = due_raw.strftime("%Y-%m-%d")
        else:
            due_str = _cell_str(row, due_idx)
        candidates.append({
            "title": title,
            "type": type_str,
            "tier": tier_str,
            "due_date": due_str,
            "linked_goal": _cell_str(row, col_map.get("linked_goal")),
            "linked_project": _cell_str(row, col_map.get("linked_project")),
            "notes": _cell_str(row, col_map.get("notes")),
            "url": _cell_str(row, col_map.get("url")),
            "included": True,
        })
    return candidates


def parse_excel_projects(file_bytes: bytes) -> list[dict[str, Any]]:
    """#80 (2026-04-26): parse a .xlsx of project rows.

    Expected columns (case-insensitive header row):
      - name (required)
      - type: work, personal
      - target_quarter: free string e.g. "2026-Q4"
      - status: not_started, in_progress, done, on_hold
      - color: hex e.g. "#2563eb" — if absent, per-type default fills (#66)
      - actions, notes: free text
      - linked_goal: free-text title; matched case-insensitively against
        existing goals at create time. No exact-id requirement.

    Rows with no name are skipped.
    """
    import openpyxl

    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True)
    except Exception as e:
        raise ValueError(f"Cannot read Excel file: {e}") from e

    ws = wb.active
    if ws is None:
        raise ValueError("Excel file has no active worksheet")

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    headers = [str(h).strip().lower() if h else "" for h in rows[0]]
    col_map = {name: idx for idx, name in enumerate(headers) if name}

    candidates: list[dict[str, Any]] = []
    for row in rows[1:]:
        name = _cell_str(row, col_map.get("name"))
        if not name:
            continue
        type_str = _cell_str(row, col_map.get("type")).lower()
        if not _valid_enum(ProjectType, type_str):
            type_str = "work"
        status_str = _cell_str(row, col_map.get("status")).lower()
        if not _valid_enum(ProjectStatus, status_str):
            status_str = "not_started"
        candidates.append({
            "name": name,
            "type": type_str,
            "target_quarter": _cell_str(row, col_map.get("target_quarter")),
            "status": status_str,
            "color": _cell_str(row, col_map.get("color")),
            "actions": _cell_str(row, col_map.get("actions")),
            "notes": _cell_str(row, col_map.get("notes")),
            "linked_goal": _cell_str(row, col_map.get("linked_goal")),
            "included": True,
        })
    return candidates


def find_duplicate_projects(names: list[str]) -> list[str]:
    """#80: return project names that already exist (case-insensitive)."""
    if not names:
        return []
    existing = db.session.scalars(
        select(Project.name).where(Project.is_active.is_(True))
    ).all()
    existing_lower = {n.lower() for n in existing}
    return [n for n in names if n.lower() in existing_lower]


def find_duplicate_goals(titles: list[str]) -> list[str]:
    """Return titles that already exist as active goals (case-insensitive)."""
    if not titles:
        return []

    existing = db.session.scalars(
        select(Goal.title).where(Goal.is_active.is_(True))
    ).all()
    existing_lower = {t.lower() for t in existing}

    return [t for t in titles if t.lower() in existing_lower]


# --- Create records from confirmed candidates --------------------------------


def create_tasks_from_import(
    candidates: list[dict[str, Any]],
    source: str,
) -> list[Task]:
    """Create Task records from confirmed import candidates.

    Only candidates with included=True are created.
    All imported tasks land in the Inbox tier.
    Logs the import to ImportLog and stamps a shared ``batch_id`` on
    every created Task so the whole import can be undone as a group
    via the recycle bin flow.

    Args:
        candidates: List of candidate dicts from the preview.
        source: Import source identifier for the log.

    Returns:
        List of newly created Task records.
    """
    # #89 (2026-04-26): pre-load goal + project lookups once so any
    # candidate carrying `linked_goal` or `linked_project` (Excel paths)
    # resolves O(1) instead of N queries. Lower-cased title/name is the
    # match key (matches the project import behavior in #80).
    goal_index: dict[str, Any] = {}
    for g in db.session.scalars(
        select(Goal).where(Goal.is_active.is_(True))
    ):
        goal_index[g.title.strip().lower()] = g.id
    project_index: dict[str, Any] = {}
    for p in db.session.scalars(
        select(Project).where(Project.is_active.is_(True))
    ):
        project_index[p.name.strip().lower()] = p.id

    batch_id = uuid.uuid4()
    created = []
    for candidate in candidates:
        if not candidate.get("included", True):
            continue

        title = (candidate.get("title") or "").strip()
        if not title:
            continue

        task_type_str = candidate.get("type", "work")
        try:
            task_type = TaskType(task_type_str)
        except ValueError:
            task_type = TaskType.WORK

        # #76 (2026-04-25): preview row now exposes more fields. Default
        # tier remains Inbox; everything else is optional. Bad values are
        # silently coerced to defaults rather than aborting the whole
        # import — the user just edited each row, so we trust them.
        tier_str = candidate.get("tier") or "inbox"
        try:
            tier = Tier(tier_str)
        except ValueError:
            tier = Tier.INBOX

        due_date_val: date | None = None
        due_raw = (candidate.get("due_date") or "").strip()
        if due_raw:
            try:
                due_date_val = date.fromisoformat(due_raw)
            except ValueError:
                due_date_val = None

        # goal_id can be a UUID string OR fall through to a free-text
        # `linked_goal` lookup (#89 Excel path).
        goal_id_val: uuid.UUID | None = None
        goal_raw = (candidate.get("goal_id") or "").strip()
        if goal_raw:
            try:
                goal_id_val = uuid.UUID(goal_raw)
            except ValueError:
                goal_id_val = None
        if goal_id_val is None:
            linked_goal_raw = (candidate.get("linked_goal") or "").strip()
            if linked_goal_raw:
                goal_id_val = goal_index.get(linked_goal_raw.lower())

        project_id_val: uuid.UUID | None = None
        project_raw = (candidate.get("project_id") or "").strip()
        if project_raw:
            try:
                project_id_val = uuid.UUID(project_raw)
            except ValueError:
                project_id_val = None
        if project_id_val is None:
            linked_project_raw = (candidate.get("linked_project") or "").strip()
            if linked_project_raw:
                project_id_val = project_index.get(linked_project_raw.lower())

        notes = (candidate.get("notes") or "").strip() or None
        url = (candidate.get("url") or "").strip() or None

        task = Task(
            title=title,
            type=task_type,
            tier=tier,
            due_date=due_date_val,
            goal_id=goal_id_val,
            project_id=project_id_val,
            notes=notes,
            url=url,
            batch_id=batch_id,
        )
        db.session.add(task)
        created.append(task)

    if created:
        log = ImportLog(
            source=source, task_count=len(created), batch_id=batch_id
        )
        db.session.add(log)
        db.session.commit()

    return created


def create_goals_from_import(
    candidates: list[dict[str, Any]],
    source: str,
) -> list[Goal]:
    """Create Goal records from confirmed import candidates.

    Only candidates with included=True are created.
    Logs the import to ImportLog and stamps a shared ``batch_id`` on
    every created Goal so the whole import can be undone as a group
    via the recycle bin flow.

    Args:
        candidates: List of goal candidate dicts from the preview.
        source: Import source identifier for the log.

    Returns:
        List of newly created Goal records.
    """
    batch_id = uuid.uuid4()
    created = []
    for candidate in candidates:
        if not candidate.get("included", True):
            continue

        title = (candidate.get("title") or "").strip()
        if not title:
            continue

        try:
            category = GoalCategory(candidate.get("category", "work"))
        except ValueError:
            category = GoalCategory.WORK

        try:
            priority = GoalPriority(candidate.get("priority", "should"))
        except ValueError:
            priority = GoalPriority.SHOULD

        try:
            status = GoalStatus(candidate.get("status", "not_started"))
        except ValueError:
            status = GoalStatus.NOT_STARTED

        goal = Goal(
            title=title,
            category=category,
            priority=priority,
            actions=candidate.get("actions") or None,
            target_quarter=candidate.get("target_quarter") or None,
            status=status,
            notes=candidate.get("notes") or None,
            batch_id=batch_id,
        )
        db.session.add(goal)
        created.append(goal)

    if created:
        log = ImportLog(
            source=source, task_count=len(created), batch_id=batch_id
        )
        db.session.add(log)
        db.session.commit()

    return created


def create_projects_from_import(
    candidates: list[dict[str, Any]],
    source: str,
) -> list[Project]:
    """#80 (2026-04-26): create Project records from confirmed candidates.

    Mirrors create_goals_from_import: only included candidates create
    rows, batch_id stamps every row for recycle-bin batch undo, ImportLog
    captures the import. linked_goal is matched case-insensitively to an
    existing active goal title; misses are silently skipped (project
    still created with goal_id=None).
    """
    from project_service import _default_color_for_type

    if not candidates:
        return []

    # Pre-load active goals into a case-insensitive title -> id lookup
    # so per-row linked_goal resolution is O(1) instead of N queries.
    goal_index: dict[str, Any] = {}
    for g in db.session.scalars(
        select(Goal).where(Goal.is_active.is_(True))
    ):
        goal_index[g.title.strip().lower()] = g.id

    batch_id = uuid.uuid4()
    created: list[Project] = []
    for candidate in candidates:
        if not candidate.get("included", True):
            continue
        name = (candidate.get("name") or "").strip()
        if not name:
            continue

        type_str = candidate.get("type", "work")
        try:
            project_type = ProjectType(type_str)
        except ValueError:
            project_type = ProjectType.WORK

        status_str = candidate.get("status", "not_started")
        try:
            status = ProjectStatus(status_str)
        except ValueError:
            status = ProjectStatus.NOT_STARTED

        # PR28 audit fix #3: validate hex format same as create_project.
        # Bad/garbage color in an Excel cell now silently falls back to
        # the per-type default rather than persisting an injection vector.
        from project_service import _parse_color
        try:
            color = _parse_color(candidate.get("color"))
        except Exception:
            color = None
        if not color:
            color = _default_color_for_type(project_type)

        goal_id = None
        linked_raw = (candidate.get("linked_goal") or "").strip()
        if linked_raw:
            goal_id = goal_index.get(linked_raw.lower())

        project = Project(
            name=name,
            type=project_type,
            color=color,
            target_quarter=(candidate.get("target_quarter") or "").strip() or None,
            actions=(candidate.get("actions") or "").strip() or None,
            notes=(candidate.get("notes") or "").strip() or None,
            status=status,
            goal_id=goal_id,
            # PR66 audit fix #131: stamp batch_id so recycle_service can
            # find this project later for undo/restore/purge.
            batch_id=batch_id,
        )
        db.session.add(project)
        created.append(project)

    if created:
        log = ImportLog(
            source=source, task_count=len(created), batch_id=batch_id
        )
        db.session.add(log)
        db.session.commit()

    return created
