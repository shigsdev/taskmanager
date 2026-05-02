"""Integration tests for import feature (Step 17).

Two import flows:
1. OneNote text → parse into task candidates → review → create in Inbox
2. Excel .xlsx → parse into goal candidates → review → create goals

Key testing concepts:
- **Text parsing** — bullet/numbered/checkbox detection, header skipping
- **Excel parsing** — column mapping, enum validation, malformed handling
- **Duplicate detection** — existing tasks/goals flagged before import
- **ImportLog** — audit trail entry created on every import
"""
from __future__ import annotations

import io

import pytest

import auth
from models import Goal, GoalCategory, GoalPriority, ImportLog, Task, TaskType, Tier, db

# --- Helper: create Excel bytes in memory ------------------------------------


def _make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Build a .docx file in memory from paragraphs and optional table rows."""
    import docx

    doc = docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        for i, row in enumerate(table_rows):
            for j, cell_text in enumerate(row):
                table.rows[i].cells[j].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _make_xlsx(rows: list[list]) -> bytes:
    """Build an .xlsx file in memory from a list of rows.

    First row is the header. Returns raw bytes suitable for upload.
    """
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# --- OneNote text parsing (internal helpers) ---------------------------------


class TestParseOnenoteText:
    """Verify parse_onenote_text handles various OneNote formats."""

    def test_bullet_points(self):
        from import_service import parse_onenote_text

        text = "- Buy groceries\n- Call dentist\n- Review report"
        result = parse_onenote_text(text)
        assert len(result) == 3
        assert result[0]["title"] == "Buy groceries"
        assert result[1]["title"] == "Call dentist"

    def test_numbered_list(self):
        from import_service import parse_onenote_text

        text = "1. First task\n2. Second task\n3) Third task"
        result = parse_onenote_text(text)
        assert len(result) == 3
        assert result[0]["title"] == "First task"
        assert result[2]["title"] == "Third task"

    def test_checkbox_items(self):
        from import_service import parse_onenote_text

        text = "\u2610 Unchecked item\n[ ] Another unchecked\n[x] Checked item"
        result = parse_onenote_text(text)
        assert len(result) == 3
        assert result[0]["title"] == "Unchecked item"

    def test_asterisk_bullets(self):
        from import_service import parse_onenote_text

        text = "* Star bullet one\n* Star bullet two"
        result = parse_onenote_text(text)
        assert len(result) == 2
        assert result[0]["title"] == "Star bullet one"

    def test_skips_empty_lines(self):
        from import_service import parse_onenote_text

        text = "- Task one\n\n\n- Task two\n   \n- Task three"
        result = parse_onenote_text(text)
        assert len(result) == 3

    def test_skips_header_lines(self):
        from import_service import parse_onenote_text

        text = "MEETING NOTES\n- Buy groceries\n2026-04-06\n- Call dentist"
        result = parse_onenote_text(text)
        assert len(result) == 2
        assert result[0]["title"] == "Buy groceries"

    def test_skips_date_headers(self):
        from import_service import parse_onenote_text

        text = "April 6, 2026\n- Task after date"
        result = parse_onenote_text(text)
        assert len(result) == 1
        assert result[0]["title"] == "Task after date"

    def test_deduplicates_within_paste(self):
        from import_service import parse_onenote_text

        text = "- Buy milk\n- Buy Milk\n- buy milk"
        result = parse_onenote_text(text)
        assert len(result) == 1

    def test_empty_text_returns_empty(self):
        from import_service import parse_onenote_text

        assert parse_onenote_text("") == []
        assert parse_onenote_text("   \n  ") == []

    def test_plain_lines_without_bullets(self):
        from import_service import parse_onenote_text

        text = "Schedule meeting with team\nReview Q2 numbers"
        result = parse_onenote_text(text)
        assert len(result) == 2

    def test_all_candidates_default_to_work(self):
        from import_service import parse_onenote_text

        result = parse_onenote_text("- A task")
        assert result[0]["type"] == "work"
        assert result[0]["included"] is True

    def test_skips_very_short_lines(self):
        from import_service import parse_onenote_text

        text = "- x\n- Real task here"
        result = parse_onenote_text(text)
        assert len(result) == 1
        assert result[0]["title"] == "Real task here"

    def test_mixed_bullet_styles(self):
        from import_service import parse_onenote_text

        text = "- Dash bullet\n* Star bullet\n1. Numbered\n\u2610 Checkbox"
        result = parse_onenote_text(text)
        assert len(result) == 4


# --- OneNote .docx parsing ---------------------------------------------------


class TestParseOnenoteDocx:
    """Verify parse_onenote_docx handles Word documents from OneNote export."""

    def test_basic_paragraphs(self):
        from import_service import parse_onenote_docx

        docx_bytes = _make_docx(["- Buy groceries", "- Call dentist"])
        result = parse_onenote_docx(docx_bytes)
        assert len(result) == 2
        assert result[0]["title"] == "Buy groceries"
        assert result[1]["title"] == "Call dentist"

    def test_extracts_table_text(self):
        from import_service import parse_onenote_docx

        docx_bytes = _make_docx([], table_rows=[["Task from table"]])
        result = parse_onenote_docx(docx_bytes)
        assert len(result) == 1
        assert result[0]["title"] == "Task from table"

    def test_combines_paragraphs_and_tables(self):
        from import_service import parse_onenote_docx

        docx_bytes = _make_docx(
            ["- Paragraph task"],
            table_rows=[["Table task"]],
        )
        result = parse_onenote_docx(docx_bytes)
        assert len(result) == 2

    def test_skips_headers_in_docx(self):
        from import_service import parse_onenote_docx

        docx_bytes = _make_docx(["MEETING NOTES", "- Real task"])
        result = parse_onenote_docx(docx_bytes)
        assert len(result) == 1
        assert result[0]["title"] == "Real task"

    def test_deduplicates_in_docx(self):
        from import_service import parse_onenote_docx

        docx_bytes = _make_docx(["- Same task", "- Same Task", "- same task"])
        result = parse_onenote_docx(docx_bytes)
        assert len(result) == 1

    def test_empty_doc_returns_empty(self):
        from import_service import parse_onenote_docx

        docx_bytes = _make_docx([])
        result = parse_onenote_docx(docx_bytes)
        assert result == []

    def test_invalid_docx_raises(self):
        from import_service import parse_onenote_docx

        with pytest.raises(ValueError, match="Cannot read"):
            parse_onenote_docx(b"not a docx file")

    def test_candidates_default_to_work(self):
        from import_service import parse_onenote_docx

        docx_bytes = _make_docx(["- A task from docx"])
        result = parse_onenote_docx(docx_bytes)
        assert result[0]["type"] == "work"
        assert result[0]["included"] is True


# --- Excel goals parsing ----------------------------------------------------


class TestParseExcelTasks:
    """#89 (2026-04-26): Excel upload for tasks."""

    def test_basic_tasks(self):
        from import_service import parse_excel_tasks
        xlsx = _make_xlsx([
            ["title", "type", "tier"],
            ["Buy bread", "personal", "today"],
            ["Ship deck", "work", "this_week"],
        ])
        result = parse_excel_tasks(xlsx)
        assert len(result) == 2
        assert result[0]["title"] == "Buy bread"
        assert result[0]["tier"] == "today"
        assert result[1]["type"] == "work"

    def test_skips_no_title_rows(self):
        from import_service import parse_excel_tasks
        xlsx = _make_xlsx([
            ["title", "type"],
            ["Real task", "work"],
            [None, "work"],
            ["", "work"],
        ])
        result = parse_excel_tasks(xlsx)
        assert len(result) == 1

    def test_invalid_enum_falls_back_to_default(self):
        from import_service import parse_excel_tasks
        xlsx = _make_xlsx([
            ["title", "type", "tier"],
            ["X", "garbage", "alsobad"],
        ])
        result = parse_excel_tasks(xlsx)
        assert result[0]["type"] == "work"
        assert result[0]["tier"] == "inbox"

    def test_due_date_normalises_datetime_cells(self):
        """PR24 TD-1: openpyxl returns a Python datetime for date-formatted
        cells — parse_excel_tasks must call .strftime to coerce to
        YYYY-MM-DD. Previous tests only covered the string-cell branch."""
        from datetime import datetime

        from import_service import parse_excel_tasks
        xlsx = _make_xlsx([
            ["title", "due_date"],
            # Real datetime → exercises the hasattr(due_raw, "strftime") branch.
            ["Real date cell", datetime(2026, 5, 15, 9, 30)],
            # Plain string → exercises the else branch.
            ["String date cell", "2026-06-01"],
        ])
        result = parse_excel_tasks(xlsx)
        assert result[0]["due_date"] == "2026-05-15"
        assert result[1]["due_date"] == "2026-06-01"

    def test_linked_goal_resolves_at_create(self, app):
        """Excel candidate with linked_goal resolves to existing goal id."""
        from import_service import (
            create_tasks_from_import,
            parse_excel_tasks,
        )
        with app.app_context():
            g = Goal(
                title="Ship Q4",
                category=GoalCategory.WORK,
                priority=GoalPriority.MUST,
            )
            db.session.add(g)
            db.session.commit()
            goal_id = g.id
            xlsx = _make_xlsx([
                ["title", "linked_goal"],
                ["Plan kickoff", "ship q4"],
            ])
            candidates = parse_excel_tasks(xlsx)
            assert candidates[0]["linked_goal"] == "ship q4"
            tasks = create_tasks_from_import(candidates, source="test")
            assert tasks[0].goal_id == goal_id


class TestParseExcelGoals:
    """Verify parse_excel_goals handles various Excel formats."""

    def test_basic_goals(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([
            ["title", "category", "priority"],
            ["Run a marathon", "health", "must"],
            ["Learn Spanish", "personal_growth", "should"],
        ])
        result = parse_excel_goals(xlsx)
        assert len(result) == 2
        assert result[0]["title"] == "Run a marathon"
        assert result[0]["category"] == "health"
        assert result[1]["priority"] == "should"

    def test_all_columns(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([
            ["title", "category", "priority", "actions", "target_quarter", "status", "notes"],
            ["Goal A", "work", "must", "Do things", "Q2 2026", "in_progress", "Some notes"],
        ])
        result = parse_excel_goals(xlsx)
        assert len(result) == 1
        assert result[0]["actions"] == "Do things"
        assert result[0]["target_quarter"] == "Q2 2026"
        assert result[0]["status"] == "in_progress"
        assert result[0]["notes"] == "Some notes"

    def test_missing_title_column_raises(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([
            ["name", "category"],
            ["Goal", "work"],
        ])
        with pytest.raises(ValueError, match="title"):
            parse_excel_goals(xlsx)

    def test_skips_empty_title_rows(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([
            ["title", "category"],
            ["Valid goal", "work"],
            ["", "health"],
            [None, "work"],
        ])
        result = parse_excel_goals(xlsx)
        assert len(result) == 1

    def test_invalid_category_defaults_to_work(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([
            ["title", "category"],
            ["Goal", "invalid_cat"],
        ])
        result = parse_excel_goals(xlsx)
        assert result[0]["category"] == "work"

    def test_invalid_priority_defaults_to_should(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([
            ["title", "priority"],
            ["Goal", "extreme"],
        ])
        result = parse_excel_goals(xlsx)
        assert result[0]["priority"] == "should"

    def test_invalid_status_defaults_to_not_started(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([
            ["title", "status"],
            ["Goal", "unknown_status"],
        ])
        result = parse_excel_goals(xlsx)
        assert result[0]["status"] == "not_started"

    def test_deduplicates_within_file(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([
            ["title"],
            ["Same Goal"],
            ["Same Goal"],
            ["same goal"],
        ])
        result = parse_excel_goals(xlsx)
        assert len(result) == 1

    def test_empty_file_returns_empty(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([])
        result = parse_excel_goals(xlsx)
        assert result == []

    def test_header_only_returns_empty(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([["title", "category"]])
        result = parse_excel_goals(xlsx)
        assert result == []

    def test_invalid_file_raises(self):
        from import_service import parse_excel_goals

        with pytest.raises(ValueError, match="Cannot read"):
            parse_excel_goals(b"not an excel file at all")

    def test_case_insensitive_headers(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([
            ["Title", "Category", "Priority"],
            ["Goal A", "health", "must"],
        ])
        result = parse_excel_goals(xlsx)
        assert len(result) == 1
        assert result[0]["title"] == "Goal A"

    def test_title_only_file(self):
        from import_service import parse_excel_goals

        xlsx = _make_xlsx([
            ["title"],
            ["Minimal goal"],
        ])
        result = parse_excel_goals(xlsx)
        assert result[0]["category"] == "work"
        assert result[0]["priority"] == "should"


# --- Duplicate detection -----------------------------------------------------


class TestDuplicateDetection:
    """Verify duplicate detection against existing DB records."""

    def test_finds_duplicate_tasks(self, app):
        from import_service import find_duplicate_tasks

        with app.app_context():
            db.session.add(Task(title="Existing task", type=TaskType.WORK, tier=Tier.INBOX))
            db.session.commit()

            dupes = find_duplicate_tasks(["Existing task", "New task"])
            assert "Existing task" in dupes
            assert "New task" not in dupes

    def test_case_insensitive_task_duplicates(self, app):
        from import_service import find_duplicate_tasks

        with app.app_context():
            db.session.add(Task(title="Buy Milk", type=TaskType.WORK, tier=Tier.INBOX))
            db.session.commit()

            dupes = find_duplicate_tasks(["buy milk", "BUY MILK"])
            assert len(dupes) == 2

    def test_finds_duplicate_goals(self, app):
        from import_service import find_duplicate_goals

        with app.app_context():
            db.session.add(Goal(
                title="Learn Python",
                category=GoalCategory.WORK,
                priority=GoalPriority.SHOULD,
            ))
            db.session.commit()

            dupes = find_duplicate_goals(["Learn Python", "New Goal"])
            assert "Learn Python" in dupes
            assert "New Goal" not in dupes

    def test_empty_list_returns_empty(self, app):
        from import_service import find_duplicate_goals, find_duplicate_tasks

        with app.app_context():
            assert find_duplicate_tasks([]) == []
            assert find_duplicate_goals([]) == []


# --- Create tasks from import ------------------------------------------------


class TestCreateTasksFromImport:
    """Verify create_tasks_from_import creates records and logs."""

    def test_creates_included_tasks(self, app):
        from import_service import create_tasks_from_import

        with app.app_context():
            candidates = [
                {"title": "Task A", "type": "work", "included": True},
                {"title": "Task B", "type": "personal", "included": True},
            ]
            tasks = create_tasks_from_import(candidates, source="test_import")
            assert len(tasks) == 2
            assert tasks[0].title == "Task A"
            assert tasks[1].type == TaskType.PERSONAL

    def test_all_tasks_land_in_inbox(self, app):
        from import_service import create_tasks_from_import

        with app.app_context():
            tasks = create_tasks_from_import(
                [{"title": "Inbox item", "included": True}],
                source="test",
            )
            assert tasks[0].tier == Tier.INBOX

    def test_skips_excluded(self, app):
        from import_service import create_tasks_from_import

        with app.app_context():
            tasks = create_tasks_from_import(
                [
                    {"title": "Include", "included": True},
                    {"title": "Exclude", "included": False},
                ],
                source="test",
            )
            assert len(tasks) == 1

    def test_logs_import(self, app):
        from import_service import create_tasks_from_import

        with app.app_context():
            create_tasks_from_import(
                [{"title": "Logged task", "included": True}],
                source="onenote_test",
            )
            log = db.session.scalars(
                db.select(ImportLog).where(ImportLog.source == "onenote_test")
            ).first()
            assert log is not None
            assert log.task_count == 1

    def test_skips_empty_titles(self, app):
        from import_service import create_tasks_from_import

        with app.app_context():
            tasks = create_tasks_from_import(
                [{"title": "", "included": True}, {"title": "  ", "included": True}],
                source="test",
            )
            assert len(tasks) == 0

    def test_accepts_full_field_set_per_candidate(self, app):
        """#76: each candidate carries optional tier/due_date/goal_id/project_id/notes/url."""
        from datetime import date

        from import_service import create_tasks_from_import
        from models import Goal, GoalCategory, GoalPriority, Project

        with app.app_context():
            goal = Goal(
                title="g", category=GoalCategory.WORK, priority=GoalPriority.MUST,
            )
            project = Project(name="p")
            db.session.add_all([goal, project])
            db.session.commit()
            goal_id = str(goal.id)
            project_id = str(project.id)

            tasks = create_tasks_from_import(
                [
                    {
                        "title": "Full field task",
                        "type": "personal",
                        "tier": "today",
                        "due_date": "2026-12-31",
                        "goal_id": goal_id,
                        "project_id": project_id,
                        "notes": "imported with notes",
                        "url": "https://example.com",
                        "included": True,
                    },
                ],
                source="test",
            )
            assert len(tasks) == 1
            t = tasks[0]
            assert t.tier == Tier.TODAY
            assert t.type == TaskType.PERSONAL
            assert t.due_date == date(2026, 12, 31)
            assert str(t.goal_id) == goal_id
            assert str(t.project_id) == project_id
            assert t.notes == "imported with notes"
            assert t.url == "https://example.com"

    def test_invalid_optional_fields_silently_default(self, app):
        """Bad tier/due_date/goal_id are coerced to defaults (don't abort import)."""
        from import_service import create_tasks_from_import

        with app.app_context():
            tasks = create_tasks_from_import(
                [
                    {
                        "title": "Bad fields",
                        "tier": "not-a-tier",
                        "due_date": "not-a-date",
                        "goal_id": "not-a-uuid",
                        "project_id": "not-a-uuid",
                        "included": True,
                    },
                ],
                source="test",
            )
            assert len(tasks) == 1
            assert tasks[0].tier == Tier.INBOX
            assert tasks[0].due_date is None
            assert tasks[0].goal_id is None
            assert tasks[0].project_id is None


# --- Create goals from import ------------------------------------------------


class TestCreateGoalsFromImport:
    """Verify create_goals_from_import creates records and logs."""

    def test_creates_included_goals(self, app):
        from import_service import create_goals_from_import

        with app.app_context():
            candidates = [
                {"title": "Goal A", "category": "health", "priority": "must", "included": True},
                {"title": "Goal B", "category": "work", "priority": "should", "included": True},
            ]
            goals = create_goals_from_import(candidates, source="excel_test")
            assert len(goals) == 2
            assert goals[0].title == "Goal A"
            assert goals[0].category == GoalCategory.HEALTH

    def test_skips_excluded(self, app):
        from import_service import create_goals_from_import

        with app.app_context():
            goals = create_goals_from_import(
                [
                    {"title": "Include", "category": "work", "priority": "must", "included": True},
                    {"title": "Exclude", "category": "work", "priority": "must", "included": False},
                ],
                source="test",
            )
            assert len(goals) == 1

    def test_logs_import(self, app):
        from import_service import create_goals_from_import

        with app.app_context():
            create_goals_from_import(
                [{"title": "Logged goal", "category": "work",
                  "priority": "must", "included": True}],
                source="excel_goals_test",
            )
            log = db.session.scalars(
                db.select(ImportLog).where(ImportLog.source == "excel_goals_test")
            ).first()
            assert log is not None
            assert log.task_count == 1

    def test_invalid_category_defaults_to_work(self, app):
        from import_service import create_goals_from_import

        with app.app_context():
            goals = create_goals_from_import(
                [{"title": "Bad cat", "category": "invalid", "priority": "must", "included": True}],
                source="test",
            )
            assert goals[0].category == GoalCategory.WORK

    def test_invalid_priority_defaults_to_should(self, app):
        from import_service import create_goals_from_import

        with app.app_context():
            goals = create_goals_from_import(
                [{"title": "Bad pri", "category": "work", "priority": "extreme", "included": True}],
                source="test",
            )
            assert goals[0].priority == GoalPriority.SHOULD


# --- Tasks parse API endpoint ------------------------------------------------


class TestTasksParseAPI:
    """Verify POST /api/import/tasks/parse."""

    def test_parse_returns_candidates(self, authed_client):
        resp = authed_client.post(
            "/api/import/tasks/parse",
            json={"text": "- Buy groceries\n- Call dentist"},
        )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["total"] == 2
        assert body["candidates"][0]["title"] == "Buy groceries"

    def test_no_json_returns_400(self, authed_client):
        resp = authed_client.post(
            "/api/import/tasks/parse",
            data="not json",
            content_type="text/plain",
        )
        assert resp.status_code == 400

    def test_empty_text_returns_400(self, authed_client):
        resp = authed_client.post(
            "/api/import/tasks/parse",
            json={"text": ""},
        )
        assert resp.status_code == 400

    def test_flags_duplicates(self, authed_client, app):
        with app.app_context():
            db.session.add(Task(title="Existing", type=TaskType.WORK, tier=Tier.INBOX))
            db.session.commit()

        resp = authed_client.post(
            "/api/import/tasks/parse",
            json={"text": "- Existing\n- Brand new task"},
        )
        body = resp.get_json()
        dupes = [c for c in body["candidates"] if c["duplicate"]]
        non_dupes = [c for c in body["candidates"] if not c["duplicate"]]
        assert len(dupes) == 1
        assert dupes[0]["title"] == "Existing"
        assert len(non_dupes) == 1


# --- Tasks upload (.docx) API endpoint ---------------------------------------


class TestTasksUploadAPI:
    """Verify POST /api/import/tasks/upload."""

    def test_upload_returns_candidates(self, authed_client):
        docx_bytes = _make_docx(["- Buy groceries", "- Call dentist"])
        data = {
            "file": (
                io.BytesIO(docx_bytes),
                "notes.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        resp = authed_client.post(
            "/api/import/tasks/upload",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["total"] == 2
        assert body["candidates"][0]["title"] == "Buy groceries"

    def test_no_file_returns_400(self, authed_client):
        resp = authed_client.post("/api/import/tasks/upload")
        assert resp.status_code == 400

    def test_wrong_extension_returns_422(self, authed_client):
        data = {"file": (io.BytesIO(b"data"), "notes.txt", "text/plain")}
        resp = authed_client.post(
            "/api/import/tasks/upload",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 422

    def test_empty_file_returns_400(self, authed_client):
        data = {
            "file": (
                io.BytesIO(b""),
                "empty.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        resp = authed_client.post(
            "/api/import/tasks/upload",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 400

    def test_invalid_docx_returns_422(self, authed_client):
        data = {
            "file": (
                io.BytesIO(b"not a docx"),
                "bad.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        resp = authed_client.post(
            "/api/import/tasks/upload",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 422

    def test_flags_duplicate_tasks(self, authed_client, app):
        with app.app_context():
            db.session.add(Task(title="Existing", type=TaskType.WORK, tier=Tier.INBOX))
            db.session.commit()

        docx_bytes = _make_docx(["- Existing", "- Brand new task"])
        data = {
            "file": (
                io.BytesIO(docx_bytes),
                "notes.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        resp = authed_client.post(
            "/api/import/tasks/upload",
            data=data,
            content_type="multipart/form-data",
        )
        body = resp.get_json()
        dupes = [c for c in body["candidates"] if c["duplicate"]]
        assert len(dupes) == 1
        assert dupes[0]["title"] == "Existing"


# --- Tasks confirm API endpoint ----------------------------------------------


class TestTasksConfirmAPI:
    """Verify POST /api/import/tasks/confirm."""

    def test_confirm_creates_tasks(self, authed_client):
        resp = authed_client.post(
            "/api/import/tasks/confirm",
            json={
                "candidates": [
                    {"title": "New task A", "type": "work", "included": True},
                    {"title": "New task B", "type": "personal", "included": True},
                ],
                "source": "onenote_test",
            },
        )
        assert resp.status_code == 201
        assert resp.get_json()["created"] == 2

    def test_confirm_skips_excluded(self, authed_client):
        resp = authed_client.post(
            "/api/import/tasks/confirm",
            json={
                "candidates": [
                    {"title": "Include", "included": True},
                    {"title": "Exclude", "included": False},
                ],
            },
        )
        assert resp.status_code == 201
        assert resp.get_json()["created"] == 1

    def test_confirmed_tasks_in_inbox(self, authed_client):
        authed_client.post(
            "/api/import/tasks/confirm",
            json={"candidates": [{"title": "Imported task", "included": True}]},
        )
        resp = authed_client.get("/api/tasks?tier=inbox")
        titles = [t["title"] for t in resp.get_json()]
        assert "Imported task" in titles

    def test_no_json_returns_400(self, authed_client):
        resp = authed_client.post(
            "/api/import/tasks/confirm",
            data="not json",
            content_type="text/plain",
        )
        assert resp.status_code == 400

    def test_invalid_candidates_returns_422(self, authed_client):
        resp = authed_client.post(
            "/api/import/tasks/confirm",
            json={"candidates": "not a list"},
        )
        assert resp.status_code == 422


# --- Goals parse API endpoint ------------------------------------------------


class TestGoalsParseAPI:
    """Verify POST /api/import/goals/parse."""

    def test_parse_returns_candidates(self, authed_client):
        xlsx = _make_xlsx([
            ["title", "category", "priority"],
            ["Run marathon", "health", "must"],
        ])
        data = {"file": (io.BytesIO(xlsx), "goals.xlsx",
                         "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
        resp = authed_client.post(
            "/api/import/goals/parse",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["total"] == 1
        assert body["candidates"][0]["title"] == "Run marathon"

    def test_no_file_returns_400(self, authed_client):
        resp = authed_client.post("/api/import/goals/parse")
        assert resp.status_code == 400

    def test_wrong_extension_returns_422(self, authed_client):
        data = {"file": (io.BytesIO(b"data"), "goals.csv", "text/csv")}
        resp = authed_client.post(
            "/api/import/goals/parse",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 422

    def test_empty_file_returns_400(self, authed_client):
        data = {"file": (io.BytesIO(b""), "empty.xlsx",
                         "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
        resp = authed_client.post(
            "/api/import/goals/parse",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 400

    def test_invalid_excel_returns_422(self, authed_client):
        data = {"file": (io.BytesIO(b"not excel"), "bad.xlsx",
                         "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
        resp = authed_client.post(
            "/api/import/goals/parse",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 422

    def test_flags_duplicate_goals(self, authed_client, app):
        with app.app_context():
            db.session.add(Goal(
                title="Existing Goal",
                category=GoalCategory.WORK,
                priority=GoalPriority.SHOULD,
            ))
            db.session.commit()

        xlsx = _make_xlsx([
            ["title"],
            ["Existing Goal"],
            ["New Goal"],
        ])
        data = {"file": (io.BytesIO(xlsx), "goals.xlsx",
                         "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
        resp = authed_client.post(
            "/api/import/goals/parse",
            data=data,
            content_type="multipart/form-data",
        )
        body = resp.get_json()
        dupes = [c for c in body["candidates"] if c["duplicate"]]
        assert len(dupes) == 1
        assert dupes[0]["title"] == "Existing Goal"


# --- Goals confirm API endpoint ----------------------------------------------


class TestGoalsConfirmAPI:
    """Verify POST /api/import/goals/confirm."""

    def test_confirm_creates_goals(self, authed_client):
        resp = authed_client.post(
            "/api/import/goals/confirm",
            json={
                "candidates": [
                    {"title": "Goal A", "category": "health", "priority": "must", "included": True},
                ],
                "source": "excel_test",
            },
        )
        assert resp.status_code == 201
        assert resp.get_json()["created"] == 1

    def test_confirm_skips_excluded(self, authed_client):
        resp = authed_client.post(
            "/api/import/goals/confirm",
            json={
                "candidates": [
                    {"title": "Include", "category": "work", "priority": "must", "included": True},
                    {"title": "Exclude", "category": "work", "priority": "must", "included": False},
                ],
            },
        )
        assert resp.status_code == 201
        assert resp.get_json()["created"] == 1

    def test_no_json_returns_400(self, authed_client):
        resp = authed_client.post(
            "/api/import/goals/confirm",
            data="not json",
            content_type="text/plain",
        )
        assert resp.status_code == 400

    def test_invalid_candidates_returns_422(self, authed_client):
        resp = authed_client.post(
            "/api/import/goals/confirm",
            json={"candidates": "not a list"},
        )
        assert resp.status_code == 422

    def test_empty_candidates_returns_zero(self, authed_client):
        resp = authed_client.post(
            "/api/import/goals/confirm",
            json={"candidates": []},
        )
        assert resp.status_code == 201
        assert resp.get_json()["created"] == 0


# --- Import page HTML --------------------------------------------------------


class TestImportPageView:
    """Verify the /import page renders with the expected structure."""

    def test_renders_200(self, client, monkeypatch):
        monkeypatch.setattr(auth, "get_current_user_email", lambda: "me@example.com")
        resp = client.get("/import")
        assert resp.status_code == 200

    def test_has_mode_buttons(self, client, monkeypatch):
        monkeypatch.setattr(auth, "get_current_user_email", lambda: "me@example.com")
        html = client.get("/import").data.decode()
        assert 'id="importTasksBtn"' in html
        assert 'id="importDocxBtn"' in html
        assert 'id="importGoalsBtn"' in html

    def test_has_text_input(self, client, monkeypatch):
        monkeypatch.setattr(auth, "get_current_user_email", lambda: "me@example.com")
        html = client.get("/import").data.decode()
        assert 'id="importText"' in html

    def test_has_file_input(self, client, monkeypatch):
        monkeypatch.setattr(auth, "get_current_user_email", lambda: "me@example.com")
        html = client.get("/import").data.decode()
        assert 'id="importFile"' in html
        assert ".xlsx" in html

    def test_has_docx_upload(self, client, monkeypatch):
        monkeypatch.setattr(auth, "get_current_user_email", lambda: "me@example.com")
        html = client.get("/import").data.decode()
        assert 'id="importDocxFile"' in html
        assert ".docx" in html

    def test_has_review_section(self, client, monkeypatch):
        monkeypatch.setattr(auth, "get_current_user_email", lambda: "me@example.com")
        html = client.get("/import").data.decode()
        assert 'id="importReview"' in html
        assert 'id="importCandidates"' in html

    def test_has_confirm_buttons(self, client, monkeypatch):
        monkeypatch.setattr(auth, "get_current_user_email", lambda: "me@example.com")
        html = client.get("/import").data.decode()
        assert 'id="importConfirmAll"' in html
        assert 'id="importConfirmSelected"' in html

    def test_has_confirm_summary_section(self, client, monkeypatch):
        monkeypatch.setattr(auth, "get_current_user_email", lambda: "me@example.com")
        html = client.get("/import").data.decode()
        assert 'id="importConfirm"' in html
        assert 'id="importConfirmList"' in html
        assert 'id="importFinalConfirm"' in html
        assert 'id="importGoBackBtn"' in html
        assert 'id="importCancelBtn"' in html

    def test_loads_import_js(self, client, monkeypatch):
        monkeypatch.setattr(auth, "get_current_user_email", lambda: "me@example.com")
        html = client.get("/import").data.decode()
        assert "import.js" in html

    def test_requires_auth(self, client, monkeypatch):
        monkeypatch.setattr(auth, "get_current_user_email", lambda: None)
        resp = client.get("/import")
        assert resp.status_code == 302


# --- Blueprint registration --------------------------------------------------


class TestImportBlueprint:
    """Verify the import_api blueprint is registered."""

    def test_blueprint_registered(self, app):
        assert "import_api" in app.blueprints

    def test_routes_exist(self, app):
        rules = [r.rule for r in app.url_map.iter_rules()]
        assert "/api/import/tasks/parse" in rules
        assert "/api/import/tasks/confirm" in rules
        assert "/api/import/goals/parse" in rules
        assert "/api/import/tasks/upload" in rules
        assert "/api/import/goals/confirm" in rules
        # #80
        assert "/api/import/projects/parse" in rules
        assert "/api/import/projects/upload" in rules
        assert "/api/import/projects/confirm" in rules


# --- #80: Projects bulk-upload -----------------------------------------------


class TestParseProjectNamesText:
    """One name per line; bullets stripped; dedup case-insensitive."""

    def test_basic_one_per_line(self):
        from import_service import parse_project_names_text
        result = parse_project_names_text("Roadmap\nQ3 Planning\nHiring")
        assert [c["name"] for c in result] == ["Roadmap", "Q3 Planning", "Hiring"]
        assert all(c["type"] == "work" for c in result)
        assert all(c["included"] is True for c in result)

    def test_strips_bullet_markers(self):
        from import_service import parse_project_names_text
        text = "- Portal\n* Roadmaps\n1. Q3 Planning\n[ ] Onboarding"
        names = [c["name"] for c in parse_project_names_text(text)]
        assert names == ["Portal", "Roadmaps", "Q3 Planning", "Onboarding"]

    def test_dedupes_case_insensitive(self):
        from import_service import parse_project_names_text
        result = parse_project_names_text("Roadmap\nROADMAP\nroadmap")
        assert len(result) == 1

    def test_empty_returns_empty(self):
        from import_service import parse_project_names_text
        assert parse_project_names_text("") == []
        assert parse_project_names_text("   \n  ") == []


class TestCreateProjectsFromImport:
    """Verify create_projects_from_import + linked_goal resolution."""

    def test_creates_included_projects(self, app):
        from import_service import create_projects_from_import
        from models import Project
        with app.app_context():
            projects = create_projects_from_import(
                [
                    {"name": "P1", "type": "work", "included": True},
                    {"name": "P2", "type": "personal", "included": True},
                    {"name": "Skip", "type": "work", "included": False},
                ],
                source="test",
            )
            assert len(projects) == 2
            stored = {p.name: p for p in db.session.scalars(db.select(Project))}
            assert "P1" in stored and "P2" in stored
            assert stored["P2"].type.value == "personal"
            # #66 default color filled when not provided.
            assert stored["P1"].color == "#2563eb"
            assert stored["P2"].color == "#16a34a"

    def test_linked_goal_resolves_case_insensitive(self, app):
        from import_service import create_projects_from_import
        with app.app_context():
            g = Goal(
                title="Ship Calendar",
                category=GoalCategory.WORK,
                priority=GoalPriority.MUST,
            )
            db.session.add(g)
            db.session.commit()
            goal_id = g.id
            projects = create_projects_from_import(
                [{
                    "name": "Linked", "type": "work",
                    "linked_goal": "ship calendar", "included": True,
                }],
                source="test",
            )
            assert projects[0].goal_id == goal_id

    def test_linked_goal_miss_skips_silently(self, app):
        from import_service import create_projects_from_import
        with app.app_context():
            projects = create_projects_from_import(
                [{
                    "name": "Orphan", "type": "work",
                    "linked_goal": "no such goal", "included": True,
                }],
                source="test",
            )
            assert projects[0].goal_id is None

    def test_logs_import(self, app):
        from import_service import create_projects_from_import
        with app.app_context():
            create_projects_from_import(
                [{"name": "Logged", "included": True}],
                source="projects_test_import",
            )
            log = db.session.scalars(
                db.select(ImportLog).where(ImportLog.source == "projects_test_import")
            ).first()
            assert log is not None
            assert log.task_count == 1


class TestProjectsParseEndpoint:
    """API: POST /api/import/projects/parse."""

    def test_text_parse_returns_candidates(self, authed_client):
        resp = authed_client.post(
            "/api/import/projects/parse",
            json={"text": "- Portal\n- Roadmaps"},
        )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["total"] == 2
        names = [c["name"] for c in body["candidates"]]
        assert "Portal" in names

    def test_confirm_creates_projects(self, authed_client):
        resp = authed_client.post(
            "/api/import/projects/confirm",
            json={
                "candidates": [
                    {"name": "API-created", "type": "work", "included": True},
                ],
                "source": "projects_api_test",
            },
        )
        assert resp.status_code == 201
        body = resp.get_json()
        assert body["created"] == 1
        assert body["projects"][0]["name"] == "API-created"


class TestImportTemplateDownload:
    """#91: GET /api/import/template/<kind>.xlsx serves a workbook with headers."""

    def test_tasks_template_downloads(self, authed_client):
        import io

        import openpyxl
        resp = authed_client.get("/api/import/template/tasks.xlsx")
        assert resp.status_code == 200
        assert "spreadsheetml" in resp.headers["Content-Type"]
        assert "tasks_import_template.xlsx" in resp.headers.get("Content-Disposition", "")
        wb = openpyxl.load_workbook(io.BytesIO(resp.data), read_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        assert rows[0] == (
            "title", "type", "tier", "due_date",
            "linked_goal", "linked_project", "notes", "url",
        )
        assert len(rows) >= 2  # header + at least 1 example

    def test_goals_template_downloads(self, authed_client):
        import io

        import openpyxl
        resp = authed_client.get("/api/import/template/goals.xlsx")
        assert resp.status_code == 200
        wb = openpyxl.load_workbook(io.BytesIO(resp.data), read_only=True)
        rows = list(wb.active.iter_rows(values_only=True))
        assert rows[0] == (
            "title", "category", "priority", "actions",
            "target_quarter", "status", "notes",
        )

    def test_projects_template_downloads(self, authed_client):
        import io

        import openpyxl
        resp = authed_client.get("/api/import/template/projects.xlsx")
        assert resp.status_code == 200
        wb = openpyxl.load_workbook(io.BytesIO(resp.data), read_only=True)
        rows = list(wb.active.iter_rows(values_only=True))
        assert rows[0] == (
            "name", "type", "target_quarter", "status",
            "color", "actions", "notes", "linked_goal",
        )

    def test_unknown_kind_returns_404(self, authed_client):
        resp = authed_client.get("/api/import/template/widgets.xlsx")
        assert resp.status_code == 404


# --- Meeting transcript parsing ---------------------------------------------


class TestExtractActionItemsSection:
    """Verify the pre-Claude section sniffer used by parse_transcript_text."""

    def test_finds_basic_action_items_section(self):
        from import_service import extract_action_items_section

        text = (
            "# Meeting Summary\n"
            "We discussed the launch.\n\n"
            "## Action Items\n"
            "- Email Sarah the Q3 plan\n"
            "- Schedule design review\n"
        )
        body = extract_action_items_section(text)
        assert body is not None
        assert "Email Sarah" in body
        assert "Schedule design review" in body
        # Body should not include the pre-section discussion.
        assert "Meeting Summary" not in body

    def test_recognises_aliases(self):
        from import_service import extract_action_items_section

        for header in ("Next Steps", "TODOs", "To-dos", "Follow-ups", "Action Item"):
            text = f"# {header}\n- Item one\n- Item two\n"
            body = extract_action_items_section(text)
            assert body is not None, f"Header {header!r} not recognised"
            assert "Item one" in body

    def test_stops_at_next_header(self):
        from import_service import extract_action_items_section

        text = (
            "## Action Items\n"
            "- First item\n"
            "## Decisions\n"
            "- Some decision\n"
        )
        body = extract_action_items_section(text)
        assert body is not None
        assert "First item" in body
        assert "Some decision" not in body

    def test_returns_none_when_absent(self):
        from import_service import extract_action_items_section

        text = "Just a free-form note about a meeting. No structure here."
        assert extract_action_items_section(text) is None

    def test_handles_empty_input(self):
        from import_service import extract_action_items_section

        assert extract_action_items_section("") is None
        assert extract_action_items_section("   \n  ") is None


class TestParseTranscriptText:
    """Verify parse_transcript_text — Claude action-item extraction."""

    def test_raises_without_api_key(self, app, monkeypatch):
        from import_service import parse_transcript_text

        monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
        with app.app_context(), pytest.raises(RuntimeError, match="ANTHROPIC_API_KEY"):
            parse_transcript_text("Some transcript")

    def test_empty_input_returns_empty(self, app):
        from import_service import parse_transcript_text

        with app.app_context():
            assert parse_transcript_text("") == []
            assert parse_transcript_text("   \n   ") == []

    def test_returns_normalised_candidates(self, app, monkeypatch):
        from unittest.mock import patch as _patch

        from import_service import parse_transcript_text

        monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key")
        claude_reply = [
            {"title": "Email Sarah the Q3 plan", "notes": None},
            {"title": "Schedule design review", "notes": "before EOM"},
        ]
        with (
            app.app_context(),
            _patch("import_service._call_claude_for_transcript", return_value=claude_reply),
        ):
            result = parse_transcript_text(
                "## Action Items\n- Email Sarah\n- Schedule review\n"
            )
        assert len(result) == 2
        assert result[0]["title"] == "Email Sarah the Q3 plan"
        assert result[0]["type"] == "work"
        assert result[0]["included"] is True
        assert result[1]["notes"] == "before EOM"

    def test_prefers_action_items_section(self, app, monkeypatch):
        """When an explicit Action Items section is present, Claude
        should be called with ONLY that section's body — not the whole
        transcript. Verifies the option-(a) "trust pre-extracted" path.
        """
        from unittest.mock import patch as _patch

        from import_service import parse_transcript_text

        monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key")
        captured: dict[str, str] = {}

        def fake_call(api_key, payload):
            captured["payload"] = payload
            return [{"title": "Pulled from section"}]

        text = (
            "# Notes\nBlah blah\n\n"
            "## Action Items\n- Email Sarah\n- Send invoice\n"
        )
        with (
            app.app_context(),
            _patch("import_service._call_claude_for_transcript", side_effect=fake_call),
        ):
            parse_transcript_text(text)
        assert "Email Sarah" in captured["payload"]
        # Pre-section body must NOT be included.
        assert "Blah blah" not in captured["payload"]

    def test_falls_back_to_full_text_without_section(self, app, monkeypatch):
        from unittest.mock import patch as _patch

        from import_service import parse_transcript_text

        monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key")
        captured: dict[str, str] = {}

        def fake_call(api_key, payload):
            captured["payload"] = payload
            return [{"title": "An action"}]

        text = "Just free-form meeting notes with no structure at all."
        with (
            app.app_context(),
            _patch("import_service._call_claude_for_transcript", side_effect=fake_call),
        ):
            parse_transcript_text(text)
        assert captured["payload"] == text

    def test_dedupes_and_drops_blank(self, app, monkeypatch):
        from unittest.mock import patch as _patch

        from import_service import parse_transcript_text

        monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key")
        claude_reply = [
            {"title": "Email Sarah"},
            {"title": "email sarah"},  # case-duplicate
            {"title": ""},              # blank
            {"title": "x"},             # too-short
            {"title": "Schedule review"},
            {"not_title": "junk"},       # missing title
            "string-not-dict",           # wrong type
        ]
        with (
            app.app_context(),
            _patch("import_service._call_claude_for_transcript", return_value=claude_reply),
        ):
            result = parse_transcript_text("anything")
        titles = [c["title"] for c in result]
        assert titles == ["Email Sarah", "Schedule review"]

    def test_truncates_long_titles(self, app, monkeypatch):
        from unittest.mock import patch as _patch

        from import_service import parse_transcript_text

        monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key")
        long_title = "x" * 250
        with (
            app.app_context(),
            _patch(
                "import_service._call_claude_for_transcript",
                return_value=[{"title": long_title}],
            ),
        ):
            result = parse_transcript_text("anything")
        assert len(result[0]["title"]) == 100


class TestExtractTranscriptJsonArray:
    """Verify the lenient JSON extractor handles Claude's output styles."""

    def test_direct_array(self):
        from import_service import _extract_transcript_json_array

        result = _extract_transcript_json_array('[{"title": "A"}]')
        assert result == [{"title": "A"}]

    def test_markdown_code_fence(self):
        from import_service import _extract_transcript_json_array

        text = 'Here you go:\n```json\n[{"title": "A"}]\n```'
        assert _extract_transcript_json_array(text) == [{"title": "A"}]

    def test_bracket_fallback(self):
        from import_service import _extract_transcript_json_array

        text = 'Sure! [{"title": "A"}, {"title": "B"}] all done.'
        assert _extract_transcript_json_array(text) == [{"title": "A"}, {"title": "B"}]

    def test_empty_returns_empty(self):
        from import_service import _extract_transcript_json_array

        assert _extract_transcript_json_array("") == []
        assert _extract_transcript_json_array("not json at all") == []


class TestTranscriptParseAPI:
    """Verify POST /api/import/transcript/parse."""

    def test_parse_returns_candidates(self, authed_client, monkeypatch):
        from unittest.mock import patch as _patch

        monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key")
        with _patch(
            "import_service._call_claude_for_transcript",
            return_value=[{"title": "Email Sarah"}, {"title": "Send invoice"}],
        ):
            resp = authed_client.post(
                "/api/import/transcript/parse",
                json={"text": "## Action Items\n- Email Sarah\n- Send invoice"},
            )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["total"] == 2
        assert body["candidates"][0]["title"] == "Email Sarah"
        # Duplicate flag should be present even when nothing matches.
        assert body["candidates"][0]["duplicate"] is False

    def test_no_json_returns_400(self, authed_client):
        resp = authed_client.post(
            "/api/import/transcript/parse",
            data="not json",
            content_type="text/plain",
        )
        assert resp.status_code == 400

    def test_empty_text_returns_400(self, authed_client):
        resp = authed_client.post(
            "/api/import/transcript/parse",
            json={"text": ""},
        )
        assert resp.status_code == 400

    def test_missing_key_returns_503(self, authed_client, monkeypatch):
        monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
        resp = authed_client.post(
            "/api/import/transcript/parse",
            json={"text": "Some content with potential action items."},
        )
        assert resp.status_code == 503
        assert "ANTHROPIC_API_KEY" in resp.get_json()["error"]

    def test_flags_duplicates(self, authed_client, app, monkeypatch):
        from unittest.mock import patch as _patch

        with app.app_context():
            db.session.add(Task(title="Existing", type=TaskType.WORK, tier=Tier.INBOX))
            db.session.commit()

        monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key")
        with _patch(
            "import_service._call_claude_for_transcript",
            return_value=[{"title": "Existing"}, {"title": "Brand new"}],
        ):
            resp = authed_client.post(
                "/api/import/transcript/parse",
                json={"text": "## Action Items\n- Existing\n- Brand new"},
            )
        body = resp.get_json()
        dupes = [c for c in body["candidates"] if c["duplicate"]]
        assert len(dupes) == 1
        assert dupes[0]["title"] == "Existing"


class TestTranscriptUploadAPI:
    """Verify POST /api/import/transcript/upload."""

    def test_upload_md_returns_candidates(self, authed_client, monkeypatch):
        from unittest.mock import patch as _patch

        monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key")
        md_bytes = b"## Action Items\n- Email Sarah\n- Send invoice\n"
        data = {"file": (io.BytesIO(md_bytes), "notes.md", "text/markdown")}
        with _patch(
            "import_service._call_claude_for_transcript",
            return_value=[{"title": "Email Sarah"}, {"title": "Send invoice"}],
        ):
            resp = authed_client.post(
                "/api/import/transcript/upload",
                data=data,
                content_type="multipart/form-data",
            )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["total"] == 2

    def test_upload_txt_returns_candidates(self, authed_client, monkeypatch):
        from unittest.mock import patch as _patch

        monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key")
        txt_bytes = b"Random meeting notes."
        data = {"file": (io.BytesIO(txt_bytes), "notes.txt", "text/plain")}
        with _patch(
            "import_service._call_claude_for_transcript",
            return_value=[{"title": "Do a thing"}],
        ):
            resp = authed_client.post(
                "/api/import/transcript/upload",
                data=data,
                content_type="multipart/form-data",
            )
        assert resp.status_code == 200

    def test_no_file_returns_400(self, authed_client):
        resp = authed_client.post("/api/import/transcript/upload")
        assert resp.status_code == 400

    def test_wrong_extension_returns_422(self, authed_client):
        data = {"file": (io.BytesIO(b"data"), "notes.docx", "application/octet-stream")}
        resp = authed_client.post(
            "/api/import/transcript/upload",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 422

    def test_empty_file_returns_400(self, authed_client):
        data = {"file": (io.BytesIO(b""), "notes.md", "text/markdown")}
        resp = authed_client.post(
            "/api/import/transcript/upload",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code == 400

    def test_oversize_file_returns_413(self, authed_client):
        big = b"x" * (5 * 1024 * 1024 + 1)
        data = {"file": (io.BytesIO(big), "notes.md", "text/markdown")}
        resp = authed_client.post(
            "/api/import/transcript/upload",
            data=data,
            content_type="multipart/form-data",
        )
        assert resp.status_code in (413, 400)  # framework MAX_CONTENT_LENGTH may pre-empt
